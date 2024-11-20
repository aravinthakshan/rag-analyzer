from fastapi import FastAPI, UploadFile, HTTPException, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import PyPDF2
import docx
import io
import os
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor
from langchain_core.documents import Document

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise EnvironmentError("GEMINI_API_KEY not found in .env file")

# Initialize LangChain's Gemini integration
llm = ChatGoogleGenerativeAI(model="gemini-pro", google_api_key=GEMINI_API_KEY)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# Initialize FastAPI app
app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

class TextAnalysisRequest(BaseModel):
    text: str
    query: Optional[str] = None
    category: Optional[str] = "summary"

def create_text_splitter():
    return RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
    )

def create_vector_store(texts: List[str]):
    documents = [Document(page_content=text) for text in texts]
    vector_store = FAISS.from_documents(documents, embeddings)
    return vector_store

def extract_text_from_pdf(content: bytes) -> str:
    pdf_file = io.BytesIO(content)
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(content: bytes) -> str:
    docx_file = io.BytesIO(content)
    doc = docx.Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

async def analyze_with_rag(text: str, query: str, category: str) -> str:
    try:
        # Split text into chunks
        text_splitter = create_text_splitter()
        texts = text_splitter.split_text(text)
        
        # Create vector store
        vector_store = create_vector_store(texts)
        
        # Create retriever with contextual compression
        base_retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 5})
        compressor = LLMChainExtractor.from_llm(llm)  # Using LangChain's Gemini integration
        compression_retriever = ContextualCompressionRetriever(
            base_retriever=base_retriever,
            base_compressor=compressor
        )
        
        # Retrieve relevant context
        retrieved_docs = compression_retriever.get_relevant_documents(query)
        context = "\n\n".join([doc.page_content for doc in retrieved_docs])
        
        # Generate category-specific prompts
        category_prompts = {
            "summary": "Provide a comprehensive summary of the following text, highlighting the main points and key takeaways:",
            "sentiment": "Analyze the sentiment and emotional tone of the following text, providing specific examples:",
            "keywords": "Extract and explain the key concepts and important terms from the following text:",
            "entity-recognition": "Identify and categorize important entities (people, organizations, locations, etc.) from the following text:"
        }
        
        # Combine category prompt with custom query
        base_prompt = category_prompts.get(category, category_prompts["summary"])
        full_prompt = f"{base_prompt}\n\nCustom Query: {query}\n\nContext: {context}\n\nOriginal Text: {text}"
        
        # Generate response using LangChain's Gemini integration
        response = await llm.ainvoke(full_prompt)
        return response.content
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error in RAG analysis: {str(e)}")

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.post("/upload")
async def upload_file(file: UploadFile, query: Optional[str] = None, category: Optional[str] = "summary"):
    if not file:
        raise HTTPException(status_code=400, detail="No file uploaded")

    try:
        content = await file.read()
        
        # Extract text based on file type
        if file.content_type == "application/pdf":
            text = extract_text_from_pdf(content)
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(content)
        elif file.content_type == "text/plain":
            text = content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}")
        
        # Analyze using RAG
        analysis = await analyze_with_rag(
            text,
            query or "Analyze the document and provide key insights",
            category
        )
        
        return {
            "result": analysis,
            "extracted_text": text
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@app.post("/analyze_multiple")
async def analyze_multiple_documents(
    files: List[UploadFile] = File(...), 
    query: Optional[str] = None, 
    category: Optional[str] = "summary"
):
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")

    try:
        # Extract text from all documents
        documents_text = []
        for file in files:
            content = await file.read()
            
            if file.content_type == "application/pdf":
                text = extract_text_from_pdf(content)
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(content)
            elif file.content_type == "text/plain":
                text = content.decode('utf-8')
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}")
            
            documents_text.append(text)
        
        # Combine all documents' text
        combined_text = "\n\n---Document Separator---\n\n".join(documents_text)
        
        # Analyze using RAG
        analysis = await analyze_with_rag(
            combined_text,
            query or "Comprehensively analyze these documents, highlighting key insights across them",
            category
        )
        
        return {
            "result": analysis,
            "document_count": len(files)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze_text")
async def analyze_text(request: TextAnalysisRequest):
    if not request.text.strip():
        raise HTTPException(status_code=400, detail="No text provided")
    
    try:
        analysis = await analyze_with_rag(
            request.text,
            request.query or "Analyze this text and provide key insights",
            request.category
        )
        return {"result": analysis}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))